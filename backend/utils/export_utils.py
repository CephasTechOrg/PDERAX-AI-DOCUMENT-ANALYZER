import os
import uuid
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Image
from reportlab.lib.utils import ImageReader
import requests
from io import BytesIO
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportUtils:
    @staticmethod
    def export_to_txt(analysis_result: dict, output_dir: str) -> str:
        """Export analysis results to TXT file"""
        filename = f"{str(uuid.uuid4())}.txt"
        filepath = os.path.join(output_dir, filename)
        
        content = "PDERAX - AI Document Analysis Report\n"
        content += "=" * 50 + "\n\n"
        
        content += "SUMMARY:\n"
        content += "-" * 20 + "\n"
        content += analysis_result.get('summary', 'No summary available') + "\n\n"
        
        content += "KEY INSIGHTS:\n"
        content += "-" * 20 + "\n"
        insights = analysis_result.get('insights', [])
        for i, insight in enumerate(insights, 1):
            content += f"{i}. {insight}\n"
        content += "\n"
        
        content += "QUESTIONS & ANSWERS:\n"
        content += "-" * 25 + "\n"
        qa_pairs = analysis_result.get('questions_answers', [])
        for i, qa in enumerate(qa_pairs, 1):
            content += f"Q{i}: {qa.get('question', '')}\n"
            content += f"A{i}: {qa.get('answer', '')}\n\n"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
            
        return filename

    @staticmethod
    def export_to_docx(analysis_result: dict, output_dir: str) -> str:
        """Export analysis results to DOCX file"""
        filename = f"{str(uuid.uuid4())}.docx"
        filepath = os.path.join(output_dir, filename)
        
        doc = Document()

        # Add local logo (centered and resized)
        logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../frontend/assets/images/logo.png"))
        if os.path.exists(logo_path):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.2))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        doc.add_heading('PDERAX - AI Document Analysis Report', 0)
        
        # Summary section
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(analysis_result.get('summary', 'No summary available'))
        
        # Insights section
        doc.add_heading('Key Insights', level=1)
        insights = analysis_result.get('insights', [])
        for insight in insights:
            doc.add_paragraph(insight, style='List Bullet')
        
        # Q&A section
        doc.add_heading('Questions & Answers', level=1)
        qa_pairs = analysis_result.get('questions_answers', [])
        for i, qa in enumerate(qa_pairs, 1):
            doc.add_paragraph(f"Q{i}: {qa.get('question', '')}", style='List Number')
            doc.add_paragraph(f"Answer: {qa.get('answer', '')}")
            doc.add_paragraph()  # Add space between Q&A pairs
        
        doc.save(filepath)
        return filename

    @staticmethod
    def export_to_pdf(analysis_result: dict, output_dir: str) -> str:
        """Export analysis results to PDF file"""
        filename = f"{str(uuid.uuid4())}.pdf"
        filepath = os.path.join(output_dir, filename)

        # Reduce top and bottom margins for a tighter layout
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            topMargin=36,      # reduced from default (~72)
            bottomMargin=36,   # compact layout
            leftMargin=50,
            rightMargin=50
        )

        styles = getSampleStyleSheet()
        story = []

        try:
            # Add PDERAX Logo (local)
            logo_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../frontend/assets/images/logo.png"))
            if os.path.exists(logo_path):
                logo = Image(logo_path, width=70, height=70)
                logo.hAlign = 'CENTER'
                story.append(logo)
                story.append(Spacer(1, 8))
        except Exception as e:
            print(f"Could not load logo: {e}")

        # Main Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=8,
            alignment=1,
            textColor='#4f46e5',
            fontName='Helvetica-Bold'
        )
        title = Paragraph('PDERAX', title_style)
        story.append(title)

        # Subtitle
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=15,
            alignment=1,
            textColor='#6b7280',
            fontName='Helvetica-Oblique'
        )
        subtitle = Paragraph('AI that understands your document', subtitle_style)
        story.append(subtitle)

        # Report Title
        report_title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=15,
            alignment=1,
            textColor='#1f2937'
        )
        report_title = Paragraph('Document Analysis Report', report_title_style)
        story.append(report_title)

        story.append(Spacer(1, 10))

        # Summary section
        summary_heading_style = ParagraphStyle(
            'SummaryHeading',
            parent=styles['Heading2'],
            fontSize=13,
            spaceAfter=8,
            textColor='#4f46e5',
            fontName='Helvetica-Bold',
            leftIndent=0
        )
        summary_heading = Paragraph('Executive Summary', summary_heading_style)
        story.append(summary_heading)

        summary_text_style = ParagraphStyle(
            'SummaryText',
            parent=styles['BodyText'],
            fontSize=10.5,
            spaceAfter=12,
            leading=14,
            leftIndent=8,
            textColor='#374151'
        )
        summary_text = analysis_result.get('summary', 'No summary available')
        story.append(Paragraph(summary_text, summary_text_style))
        story.append(Spacer(1, 10))

        # Insights section
        insights_heading = Paragraph('Key Insights', summary_heading_style)
        story.append(insights_heading)

        insights = analysis_result.get('insights', [])
        if insights:
            insight_items = []
            insight_style = ParagraphStyle(
                'InsightText',
                parent=styles['BodyText'],
                fontSize=10,
                leading=13,
                leftIndent=10,
                textColor='#4b5563'
            )
            for insight in insights:
                insight_items.append(ListItem(Paragraph(f"• {insight}", insight_style)))
            story.append(ListFlowable(insight_items, bulletType='bullet'))
        else:
            story.append(Paragraph('No insights available', styles['BodyText']))
        story.append(Spacer(1, 10))

        # Q&A section
        qa_heading = Paragraph('Questions & Answers', summary_heading_style)
        story.append(qa_heading)

        qa_pairs = analysis_result.get('questions_answers', [])
        if qa_pairs:
            qa_question_style = ParagraphStyle(
                'QAQuestion',
                parent=styles['Heading3'],
                fontSize=10.5,
                spaceAfter=4,
                textColor='#1f2937',
                fontName='Helvetica-Bold',
                leftIndent=10
            )

            qa_answer_style = ParagraphStyle(
                'QAAnswer',
                parent=styles['BodyText'],
                fontSize=9.5,
                spaceAfter=10,
                leading=12,
                leftIndent=20,
                textColor='#6b7280'
            )

            for i, qa in enumerate(qa_pairs, 1):
                question_text = f"Q{i}: {qa.get('question', '')}"
                answer_text = f"Answer: {qa.get('answer', '')}"
                story.append(Paragraph(question_text, qa_question_style))
                story.append(Paragraph(answer_text, qa_answer_style))
                story.append(Spacer(1, 6))
        else:
            story.append(Paragraph('No Q&A available', styles['BodyText']))

        # Footer
        story.append(Spacer(1, 15))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['BodyText'],
            fontSize=8,
            alignment=1,
            textColor='#9ca3af',
            spaceBefore=8
        )
        footer = Paragraph('Generated by PDERAX - Advanced AI Document Analysis', footer_style)
        story.append(footer)

        doc.build(story)
        return filename

    @staticmethod
    def export_analysis(analysis_result: dict, output_dir: str = "static/temp") -> dict:
        """Export analysis in multiple formats"""
        os.makedirs(output_dir, exist_ok=True)
        
        return {
            "txt": ExportUtils.export_to_txt(analysis_result, output_dir),
            "docx": ExportUtils.export_to_docx(analysis_result, output_dir),
            "pdf": ExportUtils.export_to_pdf(analysis_result, output_dir)
        }
